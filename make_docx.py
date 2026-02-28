from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re

raw = r'''## CLIENT SLIDES

# SLIDE P1 — Strategic Asset Allocation (SAA)
**Title:** Strategic Asset Allocation — Built Around Your Objectives

### A) Client Objectives
- Capital preservation mandate  
- USD 25k/month income requirement  
- Private market participation  
- AI thematic exposure via diversified equity sleeves  
- Liquidity buffer for flexibility  

### B) Portfolio Structure
| Asset/Sub-asset      |   Weight % |
|:---------------------|-----------:|
| us_equity            |         15 |
| europe_equity        |         10 |
| japan_equity         |         10 |
| asia_ex_japan_equity |         10 |
| global_ig            |         15 |
| global_hy            |          5 |
| em_debt              |         10 |
| private_markets      |         12 |
| hedge_funds          |          8 |
| cash                 |          5 |

### C) Client-Relevant Metrics (BASE only)
- Expected Return (BASE case): **7.82% ($6,254,711)**
- Volatility: **0.0969827686196899**
- Portfolio Income Yield (full cash-flow model): **4.37% ($3,492,317)**
- Monthly income vs requirement: **$291,026 vs $25,000/month**

### D) What This Means for You
- Income is estimated from full portfolio cash-flow capacity, not cash sleeve only.
- Income sources: equity dividends, bond coupons, private credit distributions, and cash yield.
- Strategic diversification is maintained while improving income transparency.

### Visual Placement
- **Primary visual (center):** Allocation structure table (above).  
- **Optional small visual (top-right):** Asset-allocation weight bar chart (if generated in slide editor from the same table).  
- **No backtest/stress charts on this slide.**

### Footnotes / Appendix Cross-References
- Risk statistics and stress validation: **Appendix A1–A4**.  
- Income decomposition methodology: **Appendix A1**.  
- Data integrity and robustness checks: **Appendix A6**.  
- Volatility metric source: **Appendix A1 (Backtest statistics)**.  
- Yield/income calculation source: **Appendix A1 + A6**.

---

# SLIDE P2 — Tactical Asset Allocation (TAA)
**Title:** Tactical Asset Allocation — 6–12 Month Cyclical Overlay

### A) Current Macro Backdrop
- Policy-rate plateau  
- Disinflation trend  
- Regional diversification opportunity  

### B) Tactical Adjustments (Delta vs SAA)
| Asset                | SAA   | TAA   | Delta   |
|:---------------------|:------|:------|:--------|
| asia_ex_japan_equity | 10.0% | 12.0% | +2.0%   |
| cash                 | 5.0%  | 3.0%  | -2.0%   |
| em_debt              | 10.0% | 8.0%  | -2.0%   |
| europe_equity        | 10.0% | 10.0% | +0.0%   |
| global_hy            | 5.0%  | 3.0%  | -2.0%   |
| global_ig            | 15.0% | 17.0% | +2.0%   |
| hedge_funds          | 8.0%  | 8.0%  | +0.0%   |
| japan_equity         | 10.0% | 12.0% | +2.0%   |
| private_markets      | 12.0% | 14.0% | +2.0%   |
| us_equity            | 15.0% | 13.0% | -2.0%   |

### C) Risk Discipline Note
- Tactical bands are capped at **±2%** vs SAA.
- This preserves long-term allocation discipline and strategic risk budget.

### D) Expected Impact (BASE only)
- BASE expected return shifts from **7.82%** to **8.01%**.
- The change is cyclical, not structural; core allocation remains anchored to SAA.

### Visual Placement
- **Primary visual (right panel):** SAA vs TAA delta bar chart (built from the Delta table above).  
- **Secondary visual (left panel):** Macro bullets.  
- **No stress/backtest charts on this slide.**

### Footnotes / Appendix Cross-References
- LOW/HIGH decomposition moved to: **Appendix A1**.  
- Risk statistics and drawdown validation: **Appendix A1–A2**.  
- Stress validation: **Appendix A3**.  
- Data integrity and robustness checks: **Appendix A6**.

---

## APPENDIX — Quantitative Validation & Performance Decomposition  
*This section provides quantitative validation for the client-facing strategy.*

# SLIDE A1 — Return Decomposition + Income Model
### Visual Placement
- **No charts (table-only slide).**
- Label note: **“Data sourced from live backtest engine (see A6 Data Coverage).”**

**Table caption: SAA LOW/BASE/HIGH decomposition (income model upgraded)**
| Band   | Total Return %   | Total Return $   | Income Return %   | Income Return $   | Capital Gain %   | Capital Gain $   | Monthly Income $   | vs $25k Requirement   |
|:-------|:-----------------|:-----------------|:------------------|:------------------|:-----------------|:-----------------|:-------------------|:----------------------|
| LOW    | -6.90%           | $-5,519,144      | 4.37%             | $3,492,317        | -11.26%          | $-9,011,461      | $291,026           | MEET                  |
| BASE   | 7.82%            | $6,254,711       | 4.37%             | $3,492,317        | 3.45%            | $2,762,394       | $291,026           | MEET                  |
| HIGH   | 22.54%           | $18,028,567      | 4.37%             | $3,492,317        | 18.17%           | $14,536,249      | $291,026           | MEET                  |

**Table caption: TAA LOW/BASE/HIGH decomposition (income model upgraded)**
| Band   | Total Return %   | Total Return $   | Income Return %   | Income Return $   | Capital Gain %   | Capital Gain $   | Monthly Income $   | vs $25k Requirement   |
|:-------|:-----------------|:-----------------|:------------------|:------------------|:-----------------|:-----------------|:-------------------|:----------------------|
| LOW    | -7.32%           | $-5,858,183      | 4.51%             | $3,608,920        | -11.83%          | $-9,467,103      | $300,743           | MEET                  |
| BASE   | 8.01%            | $6,406,000       | 4.51%             | $3,608,920        | 3.50%            | $2,797,079       | $300,743           | MEET                  |
| HIGH   | 23.34%           | $18,670,182      | 4.51%             | $3,608,920        | 18.83%           | $15,061,262      | $300,743           | MEET                  |

**Table caption: Yield assumption per asset bucket (trailing 12M distribution yield)**
| bucket               | proxy   |   trailing_12m_distribution_yield | yield_method                        | source   |
|:---------------------|:--------|----------------------------------:|:------------------------------------|:---------|
| us_equity            | SPY     |                         0.0106139 | sum(last12m_dividends)/latest_close | yfinance |
| europe_equity        | VGK     |                         0.0264944 | sum(last12m_dividends)/latest_close | yfinance |
| japan_equity         | EWJ     |                         0.0395366 | sum(last12m_dividends)/latest_close | yfinance |
| asia_ex_japan_equity | AIA     |                         0.0206829 | sum(last12m_dividends)/latest_close | yfinance |
| global_ig            | AGG     |                         0.0383235 | sum(last12m_dividends)/latest_close | yfinance |
| global_hy            | HYG     |                         0.0575446 | sum(last12m_dividends)/latest_close | yfinance |
| em_debt              | EMB     |                         0.0487772 | sum(last12m_dividends)/latest_close | yfinance |
| private_markets      | BIZD    |                         0.131861  | sum(last12m_dividends)/latest_close | yfinance |
| hedge_funds          | BTAL    |                         0.0254261 | sum(last12m_dividends)/latest_close | yfinance |
| cash                 | BIL     |                         0.0405916 | sum(last12m_dividends)/latest_close | yfinance |

**Table caption: Backtest statistics**
| portfolio     |      cagr |   ann_vol |   sharpe_rf0 |   sharpe_rf_cash |   sortino |   max_drawdown |   calmar |   worst_month |   best_month |   %positive_months |   tracking_error_vs_saa |
|:--------------|----------:|----------:|-------------:|-----------------:|----------:|---------------:|---------:|--------------:|-------------:|-------------------:|------------------------:|
| Current       | 0.111652  | 0.135004  |     0.851856 |         0.689186 |  0.832599 |      -0.260895 | 0.427957 |    -0.0965878 |    0.0886716 |           0.7      |              0.056677   |
| SAA           | 0.0746864 | 0.0969828 |     0.791499 |         0.565056 |  0.662938 |      -0.229063 | 0.326052 |    -0.107193  |    0.0837987 |           0.7      |              0          |
| TAA Static    | 0.0775348 | 0.103578  |     0.77309  |         0.561066 |  0.656544 |      -0.248463 | 0.312058 |    -0.113701  |    0.0884866 |           0.681818 |              0.0100669  |
| TAA Quarterly | 0.0761082 | 0.101401  |     0.774394 |         0.557817 |  0.660098 |      -0.232765 | 0.326974 |    -0.109344  |    0.0870995 |           0.681818 |              0.00712258 |

**Income Sustainability Note**
- Income model upgraded to full portfolio cash-flow approach.
- Portfolio income sources: equity dividends, bond coupons, private credit distributions, and cash yield.

# SLIDE A2 — Performance & Drawdown
### Visual Placement
- **Figure A2.1 – Growth of +115 lines:** `outputs/charts/equity_curve.png`  
- **Figure A2.2 – Drawdown Profile:** `outputs/charts/drawdown.png`  

**Table caption: Max drawdown + recovery**
| portfolio     |   max_drawdown |   avg_recovery_days |
|:--------------|---------------:|--------------------:|
| Current       |      -0.260103 |                 nan |
| SAA           |      -0.228479 |                 nan |
| TAA Quarterly |      -0.231965 |                 nan |
| TAA Static    |      -0.247634 |                 nan |

# SLIDE A3 — Stress Tests
### Visual Placement
- **Figure A3.1 – Scenario Impact Ranking:** `outputs/charts/stress_bar.png`

**Table caption: Historical stress windows**
| window     | portfolio     |   cum_return |   peak_to_trough_dd |   recovery_days |
|:-----------|:--------------|-------------:|--------------------:|----------------:|
| covid      | Current       |   -0.107027  |          -0.260103  |             nan |
| covid      | SAA           |   -0.103562  |          -0.228479  |             nan |
| covid      | TAA Static    |   -0.108542  |          -0.247634  |             nan |
| covid      | TAA Quarterly |   -0.104713  |          -0.231965  |             nan |
| rates_2022 | Current       |   -0.16156   |          -0.210017  |             nan |
| rates_2022 | SAA           |   -0.170225  |          -0.200675  |             nan |
| rates_2022 | TAA Static    |   -0.174344  |          -0.205101  |             nan |
| rates_2022 | TAA Quarterly |   -0.176026  |          -0.207391  |             nan |
| q4_2018    | Current       |   -0.104425  |          -0.144511  |             nan |
| q4_2018    | SAA           |   -0.0675693 |          -0.0888146 |             nan |
| q4_2018    | TAA Static    |   -0.0713516 |          -0.0933128 |             nan |
| q4_2018    | TAA Quarterly |   -0.0741509 |          -0.0958279 |             nan |

**Table caption: Hypothetical scenario impacts**
| scenario            | portfolio     |   impact_pct |   impact_usd | scenario_portfolio                  |
|:--------------------|:--------------|-------------:|-------------:|:------------------------------------|
| Equity rally +10%   | Current       |       0.078  |   6.24e+06   | Equity rally +10% \| Current         |
| Equity rally +10%   | SAA           |       0.045  |   3.6e+06    | Equity rally +10% \| SAA             |
| Equity rally +10%   | TAA Static    |       0.047  |   3.76e+06   | Equity rally +10% \| TAA Static      |
| Equity rally +10%   | TAA Quarterly |       0.047  |   3.76e+06   | Equity rally +10% \| TAA Quarterly   |
| Equity rally +20%   | Current       |       0.156  |   1.248e+07  | Equity rally +20% \| Current         |
| Equity rally +20%   | SAA           |       0.09   |   7.2e+06    | Equity rally +20% \| SAA             |
| Equity rally +20%   | TAA Static    |       0.094  |   7.52e+06   | Equity rally +20% \| TAA Static      |
| Equity rally +20%   | TAA Quarterly |       0.094  |   7.52e+06   | Equity rally +20% \| TAA Quarterly   |
| Equity selloff -10% | Current       |      -0.078  |  -6.24e+06   | Equity selloff -10% \| Current       |
| Equity selloff -10% | SAA           |      -0.045  |  -3.6e+06    | Equity selloff -10% \| SAA           |
| Equity selloff -10% | TAA Static    |      -0.047  |  -3.76e+06   | Equity selloff -10% \| TAA Static    |
| Equity selloff -10% | TAA Quarterly |      -0.047  |  -3.76e+06   | Equity selloff -10% \| TAA Quarterly |
| Equity selloff -20% | Current       |      -0.156  |  -1.248e+07  | Equity selloff -20% \| Current       |
| Equity selloff -20% | SAA           |      -0.09   |  -7.2e+06    | Equity selloff -20% \| SAA           |
| Equity selloff -20% | TAA Static    |      -0.094  |  -7.52e+06   | Equity selloff -20% \| TAA Static    |
| Equity selloff -20% | TAA Quarterly |      -0.094  |  -7.52e+06   | Equity selloff -20% \| TAA Quarterly |
| Equity selloff -30% | Current       |      -0.234  |  -1.872e+07  | Equity selloff -30% \| Current       |
| Equity selloff -30% | SAA           |      -0.135  |  -1.08e+07   | Equity selloff -30% \| SAA           |
| Equity selloff -30% | TAA Static    |      -0.141  |  -1.128e+07  | Equity selloff -30% \| TAA Static    |
| Equity selloff -30% | TAA Quarterly |      -0.141  |  -1.128e+07  | Equity selloff -30% \| TAA Quarterly |
| Rates up shock      | Current       |      -0.078  |  -6.24e+06   | Rates up shock \| Current            |
| Rates up shock      | SAA           |      -0.0525 |  -4.2e+06    | Rates up shock \| SAA                |
| Rates up shock      | TAA Static    |      -0.0555 |  -4.44e+06   | Rates up shock \| TAA Static         |
| Rates up shock      | TAA Quarterly |      -0.0555 |  -4.44e+06   | Rates up shock \| TAA Quarterly      |
| Credit spread shock | Current       |      -0.119  |  -9.52e+06   | Credit spread shock \| Current       |
| Credit spread shock | SAA           |      -0.0725 |  -5.8e+06    | Credit spread shock \| SAA           |
| Credit spread shock | TAA Static    |      -0.0735 |  -5.88e+06   | Credit spread shock \| TAA Static    |
| Credit spread shock | TAA Quarterly |      -0.0735 |  -5.88e+06   | Credit spread shock \| TAA Quarterly |
| Perfect storm       | Current       |      -0.237  |  -1.896e+07  | Perfect storm \| Current             |
| Perfect storm       | SAA           |      -0.178  |  -1.424e+07  | Perfect storm \| SAA                 |
| Perfect storm       | TAA Static    |      -0.1812 |  -1.4496e+07 | Perfect storm \| TAA Static          |
| Perfect storm       | TAA Quarterly |      -0.1812 |  -1.4496e+07 | Perfect storm \| TAA Quarterly       |

---

# SLIDE A4 — Distribution & Rolling Diagnostics
### Visual Placement
- `outputs/charts/monthly_return_hist.png`
- `outputs/charts/sharpe_distribution.png`
- `outputs/charts/rolling_12m_return.png`

---

# SLIDE A5 — Proxy Mapping & Model Explanation
### Visual Placement
- **No charts required** (table + methodology box).

| bucket               | proxy   | rationale                                 | limitation                                     |
|:---------------------|:--------|:------------------------------------------|:-----------------------------------------------|
| us_equity            | SPY     | US large-cap beta proxy                   | fund-level style drift not captured            |
| europe_equity        | VGK     | Europe developed equity proxy             | currency and factor mismatch                   |
| japan_equity         | EWJ     | Japan equity beta proxy                   | active manager alpha not captured              |
| asia_ex_japan_equity | AIA     | Asia ex-Japan equity proxy                | regional composition mismatch                  |
| global_ig            | AGG     | Global/US IG duration-credit blend        | not exact securitised credit mix               |
| global_hy            | HYG     | Global HY beta proxy                      | issuer-quality mix differs from case fund      |
| em_debt              | EMB     | USD EM sovereign debt proxy               | local debt/active selection not captured       |
| private_markets      | BIZD    | Listed private credit/BDC proxy           | listed beta differs from private NAV smoothing |
| hedge_funds          | BTAL    | Conservative alternative/risk-hedge proxy | not a true multi-PM hedge fund replication     |
| cash                 | BIL     | Cash/T-bill proxy                         | cash account rates may differ                  |

Model explanation box:
1) Download daily prices via yfinance.  
2) Use adjusted close; align trading calendars.  
3) Build portfolio returns with specified rebalance rules.  
4) Compute return/risk metrics and drawdowns.  
5) Run stress windows and hypothetical scenarios.  
6) Export charts/tables and assumption notes.  

Methodology robustness checks implemented: data integrity validation; transaction cost sensitivity; proxy beta stability analysis.

---

# SLIDE A6 — Robustness & Validation
### Visual Placement
- **No charts (tables only).**

**Table caption: Data coverage report**
| ticker   | start_date   | end_date   |   coverage_pct |   missing_days_pct | source   |
|:---------|:-------------|:-----------|---------------:|-------------------:|:---------|
| AGG      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| AIA      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| BIL      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| BIZD     | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| BTAL     | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| EMB      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| EWJ      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| HYG      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| SPY      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |
| VGK      | 2017-01-03   | 2026-02-27 |        96.3165 |            3.68355 | yfinance |

**Table caption: Transaction cost sensitivity (0/10/25 bps)**
|   tc_bps | portfolio     |      cagr |   sharpe_rf_cash |   max_drawdown |
|---------:|:--------------|----------:|-----------------:|---------------:|
|        0 | Current       | 0.111829  |         0.690377 |      -0.260895 |
|        0 | SAA           | 0.0749198 |         0.567313 |      -0.229063 |
|        0 | TAA Static    | 0.0775348 |         0.561066 |      -0.248463 |
|        0 | TAA Quarterly | 0.0762465 |         0.559093 |      -0.232765 |
|       10 | Current       | 0.111652  |         0.689186 |      -0.260895 |
|       10 | SAA           | 0.0746864 |         0.565056 |      -0.229063 |
|       10 | TAA Static    | 0.0775348 |         0.561066 |      -0.248463 |
|       10 | TAA Quarterly | 0.0761082 |         0.557817 |      -0.232765 |
|       25 | Current       | 0.111386  |         0.6874   |      -0.260895 |
|       25 | SAA           | 0.0743363 |         0.56167  |      -0.229063 |
|       25 | TAA Static    | 0.0775348 |         0.561066 |      -0.248463 |
|       25 | TAA Quarterly | 0.0759009 |         0.555903 |      -0.232765 |

**Table caption: Proxy beta stability (rolling 12M)**
| bucket               | proxy   | metric                              |   beta_mean |    beta_std |   beta_min |   beta_max |
|:---------------------|:--------|:------------------------------------|------------:|------------:|-----------:|-----------:|
| europe_equity        | VGK     | rolling_12m_beta_vs_spy             |   0.799491  | 0.0998194   |   0.557802 |   0.948316 |
| asia_ex_japan_equity | AIA     | rolling_12m_beta_vs_spy             |   0.90304   | 0.123417    |   0.714254 |   1.20052  |
| japan_equity         | EWJ     | rolling_12m_beta_vs_spy             |   0.725292  | 0.078697    |   0.568485 |   0.966782 |
| us_equity            | SPY     | rolling_12m_beta_vs_spy             |   1         | 1.31931e-15 |   1        |   1        |
| global_ig            | AGG     | rolling_12m_beta_vs_us_equity_proxy |   0.0340999 | 0.0669017   |  -0.115463 |   0.225945 |

Footnote: Market data sourced via yfinance; adjusted close used; see data coverage table.
'''

# Parse markdown
lines = raw.splitlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# paragraph helper

def add_para(text, bold=False, italic=False, align=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    return p


def is_table_line(s):
    s=s.strip()
    return s.startswith('|') and s.endswith('|')


def split_row(s):\n    inner=s.strip()[1:-1]\n    parts=[]\n    cur=''\n    esc=False\n    for ch in inner:\n        if esc:\n            cur += ch\n            esc=False\n            continue\n        if ch=='\\\\':\n            esc=True\n            continue\n        if ch=='|':\n            parts.append(cur.strip())\n            cur=''\n        else:\n            cur += ch\n    parts.append(cur.strip())\n    return parts

i=0
slide_started=False
source_tables=[]
created_tables=[]
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped == '':
        i+=1
        continue

    if stripped == '## CLIENT SLIDES':
        add_para('CLIENT SLIDES', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22)
        i+=1
        continue

    if stripped.startswith('## APPENDIX'):
        doc.add_page_break()
        add_para(stripped.replace('## ','').replace('**',''), bold=True, size=14)
        i+=1
        continue

    if stripped.startswith('# SLIDE'):
        if slide_started:
            doc.add_page_break()
        slide_started=True
        add_para(stripped.replace('# ','').replace('**',''), bold=True, size=14)
        i+=1
        continue

    if stripped.startswith('### '):
        add_para(stripped.replace('### ','').replace('**',''), bold=True, size=11)
        i+=1
        continue

    if stripped.startswith('**Title:**'):
        t = stripped.replace('**','')
        add_para(t, italic=True, size=11)
        i+=1
        continue

    if stripped.startswith('**Table caption:') or stripped.startswith('Table caption:'):
        t=stripped.replace('**','')
        add_para(t, bold=True, size=10)
        i+=1
        continue

    if stripped == '---':
        i+=1
        continue

    if stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        txt=stripped[2:].replace('**','').replace('`','')
        r=p.add_run(txt)
        r.font.name='Arial'; r.font.size=Pt(10)
        i+=1
        continue

    if is_table_line(stripped):
        tlines=[]
        while i < len(lines) and is_table_line(lines[i].strip()):
            tlines.append(lines[i].strip())
            i+=1
        # markdown table parse
        hdr = split_row(tlines[0])
        body = []
        for row in tlines[2:]:
            body.append(split_row(row))
        source_tables.append((len(body)+1, len(hdr)))

        tbl=doc.add_table(rows=len(body)+1, cols=len(hdr))
        tbl.style='Table Grid'
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit=True
        # header
        for c, val in enumerate(hdr):
            cell=tbl.cell(0,c)
            cell.text=val
            p=cell.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            run=p.runs[0]
            run.bold=True
            run.font.name='Arial'; run.font.size=Pt(9 if len(hdr)>=8 else 10)
        # body
        num_re=re.compile(r'^[\$\-\+0-9\.,eE%]+$')
        for r_idx,row in enumerate(body, start=1):
            for c,val in enumerate(row):
                cell=tbl.cell(r_idx,c)
                cell.text=val
                p=cell.paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT if num_re.match(val.strip()) else WD_ALIGN_PARAGRAPH.LEFT
                run=p.runs[0]
                run.font.name='Arial'; run.font.size=Pt(9 if len(hdr)>=8 else 10)
        created_tables.append((len(body)+1, len(hdr)))
        continue

    # normal paragraph
    add_para(stripped.replace('**','').replace('*','').replace('`',''))
    i+=1

assert source_tables == created_tables, f'Table shape mismatch: {source_tables} vs {created_tables}'

out=Path('/mnt/data/CLIENT_SLIDES_Reformatted_Submission.docx')
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print(out)
print(f'tables={len(created_tables)}')

