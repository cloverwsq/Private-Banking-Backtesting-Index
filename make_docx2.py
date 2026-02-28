from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re

src = Path('outputs/slide_content.md')
text = src.read_text(encoding='utf-8')
lines = text.splitlines()

doc=Document()
style=doc.styles['Normal']
style.font.name='Arial'; style.font.size=Pt(10)


def addp(t,b=False,i=False,size=10,align=None):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0)
    p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.line_spacing=1.1
    if align: p.alignment=align
    r=p.add_run(t)
    r.bold=b; r.italic=i; r.font.name='Arial'; r.font.size=Pt(size)
    return p

def is_table_line(s):
    s=s.strip(); return s.startswith('|') and s.endswith('|')

def split_row(s):
    inner=s.strip()[1:-1]
    out=[]; cur=''; esc=False
    for ch in inner:
        if esc:
            cur += ch; esc=False; continue
        if ch == '\\':
            esc=True; continue
        if ch == '|':
            out.append(cur.strip()); cur=''
        else:
            cur += ch
    out.append(cur.strip())
    return out

slide_started=False
src_shapes=[]; out_shapes=[]
i=0
while i < len(lines):
    raw=lines[i]; s=raw.strip()
    if not s:
        i+=1; continue
    if s.startswith('## CLIENT SLIDES'):
        addp('CLIENT SLIDES', b=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER); i+=1; continue
    if s.startswith('## APPENDIX'):
        doc.add_page_break(); addp(s.replace('## ','').replace('**','').replace('*',''), b=True, size=14); i+=1; continue
    if s.startswith('# SLIDE'):
        if slide_started: doc.add_page_break()
        slide_started=True
        addp(s.replace('# ','').replace('**',''), b=True, size=14); i+=1; continue
    if s.startswith('### '):
        addp(s.replace('### ','').replace('**',''), b=True, size=11); i+=1; continue
    if s.startswith('Title:') or s.startswith('**Title:**'):
        addp(s.replace('**',''), i=True, size=11); i+=1; continue
    if s.startswith('Table caption:') or s.startswith('**Table caption:'):
        addp(s.replace('**',''), b=True, size=10); i+=1; continue
    if s == '---':
        i+=1; continue
    if s.startswith('- '):
        p=doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.1
        r=p.add_run(s[2:].replace('**','').replace('`',''))
        r.font.name='Arial'; r.font.size=Pt(10)
        i+=1; continue
    if is_table_line(s):
        tlines=[]
        while i < len(lines) and is_table_line(lines[i].strip()):
            tlines.append(lines[i].strip()); i+=1
        hdr=split_row(tlines[0])
        body=[split_row(r) for r in tlines[2:]]
        src_shapes.append((len(body)+1,len(hdr)))
        tbl=doc.add_table(rows=len(body)+1, cols=len(hdr))
        tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=True
        for c,v in enumerate(hdr):
            cell=tbl.cell(0,c); cell.text=v
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            r=p.runs[0]; r.bold=True; r.font.name='Arial'; r.font.size=Pt(9 if len(hdr)>=8 else 10)
        num_re=re.compile(r'^[\$\-\+0-9\.,eE%]+$')
        for rr,row in enumerate(body,1):
            if len(row)!=len(hdr):
                row=row[:len(hdr)] + ['']*(len(hdr)-len(row))
            for c,v in enumerate(row):
                cell=tbl.cell(rr,c); cell.text=v
                p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT if num_re.match(v.strip()) else WD_ALIGN_PARAGRAPH.LEFT
                r=p.runs[0]; r.font.name='Arial'; r.font.size=Pt(9 if len(hdr)>=8 else 10)
        out_shapes.append((len(body)+1,len(hdr)))
        continue
    addp(s.replace('**','').replace('*','').replace('`',''))
    i+=1

assert src_shapes==out_shapes

out=Path('/mnt/data/CLIENT_SLIDES_Reformatted_Submission.docx')
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print(str(out))
print(f'tables={len(out_shapes)}')
